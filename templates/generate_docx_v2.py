# -*- coding: utf-8 -*-
"""
LaTeX → DOCX 转换脚本 v2
彻底清理所有LaTeX残留符号：{ } $ % \ & # _ ^
"""

import re, os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEX_FILE = os.path.join(os.path.dirname(__file__), '论文.tex')
FIGURES_DIR = os.path.join(os.path.dirname(__file__), 'figures')
OUTPUT_FILE = os.path.join(os.path.dirname(__file__), '论文.docx')

# Greek字母映射
GREEK_MAP = {
    'alpha': 'α', 'beta': 'β', 'gamma': 'γ', 'delta': 'δ', 'epsilon': 'ε',
    'zeta': 'ζ', 'eta': 'η', 'theta': 'θ', 'iota': 'ι', 'kappa': 'κ',
    'lambda': 'λ', 'mu': 'μ', 'nu': 'ν', 'xi': 'ξ', 'omicron': 'ο',
    'pi': 'π', 'rho': 'ρ', 'sigma': 'σ', 'tau': 'τ', 'upsilon': 'υ',
    'phi': 'φ', 'chi': 'χ', 'psi': 'ψ', 'omega': 'ω',
    'Alpha': 'Α', 'Beta': 'Β', 'Gamma': 'Γ', 'Delta': 'Δ',
    'Theta': 'Θ', 'Lambda': 'Λ', 'Pi': 'Π', 'Sigma': 'Σ',
    'Phi': 'Φ', 'Omega': 'Ω', 'varepsilon': 'ε', 'varphi': 'φ',
    'vartheta': 'ϑ', 'varrho': 'ϱ', 'varsigma': 'ς',
}
GREEK_SORTED = sorted(GREEK_MAP.items(), key=lambda x: -len(x[0]))

# LaTeX符号映射
LATEX_SYMBOLS = [
    ('\\cdot', '·'), ('\\times', '×'), ('\\approx', '≈'), ('\\sim', '~'),
    ('\\propto', '∝'), ('\\leq', '≤'), ('\\geq', '≥'), ('\\neq', '≠'),
    ('\\pm', '±'), ('\\mp', '∓'), ('\\infty', '∞'), ('\\partial', '∂'),
    ('\\nabla', '∇'), ('\\int', '∫'), ('\\sum', 'Σ'), ('\\prod', 'Π'),
    ('\\rightarrow', '→'), ('\\leftarrow', '←'), ('\\Rightarrow', '⇒'),
    ('\\Leftrightarrow', '⇔'), ('\\angle', '∠'), ('\\triangle', '△'),
    ('\\parallel', '∥'), ('\\perp', '⊥'), ('\\in', '∈'), ('\\notin', '∉'),
    ('\\subset', '⊂'), ('\\subseteq', '⊆'), ('\\forall', '∀'),
    ('\\exists', '∃'), ('\\emptyset', '∅'), ('\\mathbb{R}', 'ℝ'),
    ('\\mathbb{N}', 'ℕ'), ('\\mathbb{Z}', 'ℤ'), ('\\degree', '°'),
    ('\\quad', ' '), ('\\qquad', '  '), ('\\,', ''), ('\\;', ' '),
    ('\\textbackslash', '\\'), ('\\textasciitilde', '~'),
    ('\\textasciicircum', '^'), ('\\%', '%'), ('\\&', '&'),
    ('\\#', '#'), ('\\_', '_'), ('\\{', '{'), ('\\}', '}'),
    ('\\$', '$'), ('``', '"'), ("''", '"'),
]


# ============================================================
# 核心：LaTeX文本彻底清理函数
# ============================================================

def clean_latex_text(text):
    """
    彻底清理LaTeX语法标记，只保留可读的纯文本内容。
    这是v2的核心升级——所有文本在进入DOCX前必须经过此函数。
    """
    if not text or not text.strip():
        return ''

    # 1. 移除LaTeX注释（%开头的行）
    text = re.sub(r'(?<!\\)%.*$', '', text, flags=re.MULTILINE)

    # 2. 移除 \begin{...} 和 \end{...} 标签（整行或行内）
    text = re.sub(r'\\begin\{[^}]*\}', '', text)
    text = re.sub(r'\\end\{[^}]*\}', '', text)

    # 3. 移除 \label{...}, \ref{...}, \cite{...}, \bibitem{...}
    text = re.sub(r'\\label\{[^}]*\}', '', text)
    text = re.sub(r'\\ref\{[^}]*\}', '', text)
    text = re.sub(r'\\cite\{[^}]*\}', '', text)
    text = re.sub(r'\\bibitem\{[^}]*\}', '', text)
    text = re.sub(r'\\pageref\{[^}]*\}', '', text)

    # 4. 处理带参数的命令（保留内容，移除命令）
    # \textbf{content} → content
    # \textit{content} → content
    # \emph{content} → content
    # \texttt{content} → content
    # \text{content} → content
    # \mathbf{content} → content
    # \mathrm{content} → content
    text = re.sub(r'\\(?:textbf|textit|emph|texttt|text|mathbf|mathrm|mathit|mathsf|mathtt|boldsymbol|bm)\{([^{}]*(?:\{[^{}]*\}[^{}]*)*)\}',
                  r'\1', text)

    # 5. 处理其他 \command{arg} 形式（保留arg内容）
    text = re.sub(r'\\(?:section|subsection|subsubsection|paragraph|subparagraph|caption|footnote|hfill|hspace|vspace|vskip|hskip|rule|raisebox|makebox|framebox|parbox|mbox|fbox)\*?\{([^{}]*(?:\{[^{}]*\}[^{}]*)*)\}',
                  r'\1', text)

    # 6. 移除无参数的命令
    text = re.sub(r'\\(?:newpage|clearpage|pagebreak|linebreak|newline|centering|raggedright|raggedleft|small|large|Large|LARGE|huge|Huge|footnotesize|normalsize|tiny|scriptsize|noindent|indent|vfill|hfill|par|\\\\)\b', '', text)

    # 7. 移除 \command 形式的命令（只移除命令本身，不移除其后的文本）
    text = re.sub(r'\\(?:displaystyle|textstyle|scriptstyle|limits|nolimits|nonumber|notag|left|right|big|Big|bigg|Bigg|bigl|Bigl|biggl|Biggl|bigr|Bigr|biggr|Biggr|middle|colon|quad|qquad|enspace|enskip|thinspace|medspace|thickspace|negthinspace|negmedspace|negthickspace|allowdisplaybreaks)\b', '', text)

    # 8. 移除 \begin{tabular}{...} 和 \end{tabular}
    text = re.sub(r'\{tabular\}?\{[^}]*\}', '', text)
    text = re.sub(r'\{table\}?\{[^}]*\}', '', text)
    text = re.sub(r'\{figure\}?\{[^}]*\}', '', text)
    text = re.sub(r'\{algorithm\}?\{[^}]*\}', '', text)
    text = re.sub(r'\{enumerate\}?\{[^}]*\}', '', text)
    text = re.sub(r'\{itemize\}?\{[^}]*\}', '', text)

    # 9. 处理表格格式命令
    text = text.replace('\\toprule', '').replace('\\midrule', '').replace('\\bottomrule', '')
    text = text.replace('\\hline', '').replace('\\cline', '')
    text = re.sub(r'\\multicolumn\{[^}]*\}\{[^}]*\}\{', '', text)

    # 10. 转换LaTeX特殊字符
    text = text.replace('\\%', '%')
    text = text.replace('\\&', '&')
    text = text.replace('\\#', '#')
    text = text.replace('\\_', '_')
    text = text.replace('\\{', '{')
    text = text.replace('\\}', '}')
    text = text.replace('\\$', '$')
    text = text.replace('``', '"')
    text = text.replace("''", '"')

    # 11. 转换LaTeX符号
    for latex, symbol in LATEX_SYMBOLS:
        text = text.replace(latex, symbol)
    for greek, symbol in GREEK_SORTED:
        text = text.replace('\\' + greek, symbol)

    # 12. 处理数学模式标识符 $...$（转为可读文本）
    # 保留$中的内容但移除$符号
    text = re.sub(r'\$([^$]+)\$', r'\1', text)

    # 13. 清理 _{text} 和 ^{text} 模式（下/上标在正文中转为纯文本）
    text = re.sub(r'_\{(.*?)\}', r'\1', text)
    text = re.sub(r'\^\{(.*?)\}', r'\1', text)
    # 清理孤立的 _ 和 ^（单字符下/上标）
    text = re.sub(r'_([a-zA-Z0-9])', r'\1', text)
    text = re.sub(r'\^([a-zA-Z0-9])', r'\1', text)

    # 14. 清理孤立的 { } 括号
    for _ in range(3):
        text = re.sub(r'\{\s*\}', '', text)
        text = re.sub(r'\{([^{}]*)\}', r'\1', text)

    # 14. 移除残留的LaTeX命令（任何剩余的 \word 形式）
    text = re.sub(r'\\[a-zA-Z]+\*?(?:\{[^{}]*\})?', '', text)

    # 15. 清理多余空格
    text = re.sub(r' +', ' ', text)
    text = re.sub(r'\n\s*\n', '\n\n', text)

    # 16. 移除行首行尾的多余空格
    text = text.strip()

    return text


def clean_latex_light(text):
    """轻量清理（用于公式渲染前后的文字）"""
    text = re.sub(r'\\%', '%', text)
    text = re.sub(r'\\&', '&', text)
    text = re.sub(r'\\#', '#', text)
    text = re.sub(r'\\_', '_', text)
    text = re.sub(r'``', '"', text)
    text = re.sub(r"''", '"', text)
    for latex, symbol in LATEX_SYMBOLS:
        text = text.replace(latex, symbol)
    for greek, symbol in GREEK_SORTED:
        text = text.replace('\\' + greek, symbol)
    return text


# ============================================================
# 工具函数
# ============================================================

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            element = OxmlElement(f'w:{edge}')
            for attr in ['sz', 'val', 'color', 'space']:
                if attr in edge_data:
                    element.set(qn(f'w:{attr}'), str(edge_data[attr]))
            tcBorders.append(element)
    tcPr.append(tcBorders)


def add_run(paragraph, text, bold=False, italic=False, size=Pt(12),
            font_name=None, subscript=False, superscript=False):
    """添加黑色文字run"""
    run = paragraph.add_run(text)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = size
    if font_name:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if subscript:
        run.font.subscript = True
    if superscript:
        run.font.superscript = True
    return run


def render_math(paragraph, expr, base_size=Pt(12)):
    """渲染数学表达式：处理上下标和Greek字母"""
    # 先处理Greek字母
    for greek, symbol in GREEK_SORTED:
        expr = expr.replace('\\' + greek, symbol)

    # 处理LaTeX符号
    for latex, symbol in LATEX_SYMBOLS:
        expr = expr.replace(latex, symbol)

    # 移除常见的LaTeX格式命令
    for cmd in ['\\text', '\\mathbf', '\\boldsymbol', '\\mathrm',
                '\\mathit', '\\mathsf', '\\mathtt', '\\bm',
                '\\left', '\\right', '\\big', '\\Big', '\\bigg', '\\Bigg',
                '\\displaystyle', '\\tfrac', '\\dfrac', '\\frac',
                '\\limits', '\\nolimits', '\\nonumber', '\\notag']:
        expr = expr.replace(cmd, '')

    # 移除空的花括号组
    expr = re.sub(r'\{\s*\}', '', expr)

    i = 0
    while i < len(expr):
        if expr[i] == '_' and i + 1 < len(expr):
            i += 1
            if i < len(expr) and expr[i] == '{':
                j = expr.index('}', i)
                sub_text = expr[i+1:j]
                # 递归清理sub_text
                sub_text = clean_latex_light(sub_text)
                if sub_text:
                    add_run(paragraph, sub_text, size=Pt(base_size.pt * 0.85), subscript=True)
                i = j + 1
            elif i < len(expr):
                add_run(paragraph, expr[i], size=Pt(base_size.pt * 0.85), subscript=True)
                i += 1
        elif expr[i] == '^' and i + 1 < len(expr):
            i += 1
            if i < len(expr) and expr[i] == '{':
                j = expr.index('}', i)
                sup_text = expr[i+1:j]
                sup_text = clean_latex_light(sup_text)
                if sup_text:
                    add_run(paragraph, sup_text, size=Pt(base_size.pt * 0.85), superscript=True)
                i = j + 1
            elif i < len(expr):
                add_run(paragraph, expr[i], size=Pt(base_size.pt * 0.85), superscript=True)
                i += 1
        elif expr[i] in '{}':
            # 跳过孤立的括号
            i += 1
        elif expr[i] == '\\':
            # 跳过残留命令（应该在clean_latex_light中已处理）
            j = i + 1
            while j < len(expr) and expr[j].isalpha():
                j += 1
            if j < len(expr) and expr[j] == '{':
                k = expr.index('}', j)
                j = k + 1
            i = j
        elif expr[i] == '$':
            i += 1  # 跳过残留的$符号
        else:
            add_run(paragraph, expr[i], size=base_size)
            i += 1


def add_heading_styled(doc, text, level=1):
    """添加标题"""
    text = clean_latex_light(text)
    if level == 0:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(12)
        add_run(p, text, bold=True, size=Pt(16), font_name='黑体')
        return p
    p = doc.add_paragraph()
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, text, bold=True, size=Pt(14), font_name='黑体')
    elif level in (2, 3):
        add_run(p, text, bold=True, size=Pt(12), font_name='黑体')
    else:
        add_run(p, text, size=Pt(12))
    return p


def add_body_paragraph(doc, text, indent=True, size=Pt(12)):
    """添加正文段落"""
    text = clean_latex_text(text)
    if not text.strip():
        return None

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.3
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)

    # 如果还残留行内数学$...$，进一步清理
    text = re.sub(r'\$([^$]+)\$', r'\1', text)

    add_run(p, text, size=size, font_name='宋体')
    return p


def add_figure(doc, image_path, caption, width_inches=5.5):
    """嵌入图片"""
    full_path = os.path.join(FIGURES_DIR, image_path) if not os.path.isabs(image_path) else image_path
    if os.path.exists(full_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        try:
            run.add_picture(full_path, width=Inches(width_inches))
        except Exception as e:
            add_run(p, f'[图表: {image_path}]', italic=True, size=Pt(10))
            print(f'Warning: embed error {e}')

        caption = clean_latex_text(caption)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, caption, bold=True, size=Pt(10))
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, f'[图表缺失: {image_path}]', italic=True, size=Pt(10))
        print(f'Warning: not found {full_path}')


def add_reference(doc, ref_text):
    """参考文献条目"""
    ref_text = clean_latex_text(ref_text)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(-0.74)
    p.paragraph_format.left_indent = Cm(0.74)
    add_run(p, ref_text, size=Pt(10), font_name='宋体')


def add_code_block(doc, code_text):
    """代码块"""
    code_text = clean_latex_text(code_text)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    add_run(p, code_text, size=Pt(9), font_name='Consolas')


def add_abstract_page(doc, abstract_text, keywords_str):
    """摘要页"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, '摘  要', bold=True, size=Pt(16), font_name='黑体')
    doc.add_paragraph()

    abstract_text = clean_latex_text(abstract_text)
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_run(p, abstract_text, size=Pt(14), font_name='宋体')
    doc.add_paragraph()

    p = doc.add_paragraph()
    keywords_str = clean_latex_text(keywords_str)
    add_run(p, '关键词：', bold=True, size=Pt(14))
    add_run(p, keywords_str, size=Pt(14), font_name='宋体')
    doc.add_page_break()


# ============================================================
# 主转换
# ============================================================

def convert_latex_to_docx():
    print(f'Reading: {TEX_FILE}')
    with open(TEX_FILE, 'r', encoding='utf-8') as f:
        tex_content = f.read()

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style.font.color.rgb = RGBColor(0, 0, 0)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ---- 提取标题 ----
    # 实际LaTeX格式：{\heiti\zihao{3}\textbf{基于开普勒定律的月球绕日运动轨迹仿真模型}}
    paper_title = ''
    # 策略：查找 \textbf{...} 中的内容，优先找标题字号附近的
    title_patterns = [
        r'\{\\heiti\\zihao\{3\}\\textbf\{([^}]+)\}\}',   # 华中杯标题格式
        r'\\textbf\{([^}]+)\}',                            # 通用textbf
        r'\\title\{([^}]+)\}',                             # \title命令
    ]
    for pattern in title_patterns:
        m = re.search(pattern, tex_content)
        if m:
            paper_title = m.group(1).strip()
            break
    # 如果都没匹配到，找第一个 \textbf 作为标题
    if not paper_title:
        all_bf = re.findall(r'\\textbf\{([^}]+)\}', tex_content)
        if all_bf:
            # 找最长的那个（大概率是标题）
            paper_title = max(all_bf, key=len).strip()
    if not paper_title:
        paper_title = '论文标题'

    paper_title = clean_latex_text(paper_title)
    add_heading_styled(doc, paper_title, level=0)

    # ---- 提取摘要 ----
    abs_text = ''
    keywords = ''

    # 策略1：找到 \songti\zihao{4} 之后的正文（最可靠的摘要正文标记）
    abs_body_match = re.search(r'\\songti\\zihao\{4\}\s*\n(.*?)(?=\\vspace|\\noindent|关键词)', tex_content, re.DOTALL)
    if abs_body_match:
        abs_body = abs_body_match.group(1)
        # 清理LaTeX
        abs_body = abs_body.replace('\\%', '%').replace('\\&', '&').replace('\\#', '#')
        abs_body = re.sub(r'\{([^{}]*)\}', r'\1', abs_body)
        abs_body = re.sub(r'\$([^$]+)\$', r'\1', abs_body)
        abs_body = re.sub(r'\\[a-zA-Z]+\*?(?:\{[^{}]*\})?', '', abs_body)
        abs_body = re.sub(r'%[^\n]*', '', abs_body)
        abs_body = re.sub(r'\s+', ' ', abs_body).strip()
        abs_text = abs_body

    # 策略2：如果策略1失败，找"摘 要"标签后的文字
    if not abs_text or len(abs_text) < 20:
        abs_match = re.search(r'摘\s{1,3}要', tex_content)
        if abs_match:
            after_abs = tex_content[abs_match.end():abs_match.end() + 1000]
            # 跳过 \end{center} \vspace 等命令
            after_abs = re.sub(r'\\end\{center\}', '', after_abs)
            after_abs = re.sub(r'\\vspace\{[^}]*\}', '', after_abs)
            after_abs = re.sub(r'%[^\n]*', '', after_abs)
            after_abs = re.sub(r'\\(?:songti|heiti|zihao)\{[^}]*\}', '', after_abs)
            after_abs = re.sub(r'\{([^{}]*)\}', r'\1', after_abs)
            after_abs = re.sub(r'\$([^$]+)\$', r'\1', after_abs)
            after_abs = re.sub(r'\\[a-zA-Z]+\*?(?:\{[^{}]*\})?', '', after_abs)
            after_abs = re.sub(r'\s+', ' ', after_abs).strip()

            kw_pos = after_abs.find('关键词')
            if kw_pos > 0:
                abs_text = after_abs[:kw_pos].strip()
            else:
                abs_text = after_abs[:500]

    # 提取关键词
    kw_match = re.search(r'关键词[：:]\s*\}?\s*(.+?)(?:\}|$)', tex_content, re.DOTALL)
    if kw_match:
        keywords = kw_match.group(1).strip()
        keywords = re.sub(r'[\{\}]', '', keywords)  # 移除所有花括号
        keywords = re.sub(r'\\(?:textbf|noindent)\b', '', keywords)
        keywords = re.sub(r'[\{\}]', '', keywords)  # 再次清理
        keywords = re.sub(r'\s+', ' ', keywords).strip()
    if not abs_text or len(abs_text) < 20:
        abs_text = '摘要正文自动提取失败，请参见论文.tex原文。'
    if not keywords or len(keywords) < 5:
        # 从关键词标签周围提取
        kw_region = re.search(r'关键词[：:](.{5,100})', tex_content)
        if kw_region:
            keywords = re.sub(r'[\{\}\\]', '', kw_region.group(1)).strip()
        else:
            keywords = '开普勒定律；月球轨迹；数值仿真'

    add_abstract_page(doc, abs_text, keywords)

    # ---- 解析正文 ----
    body_start = tex_content.find('\\section{问题重述}')
    if body_start < 0:
        body_start = tex_content.find('\\section{')
    end_match = re.search(r'\\end\{document\}', tex_content)
    body_end = end_match.start() if end_match else len(tex_content)
    body = tex_content[body_start:body_end]

    # 分离参考文献
    ref_match = re.search(r'\\begin\{thebibliography\}', body)
    appendix_match = re.search(r'\\section\*\{附录\}', body)

    if ref_match:
        main_body = body[:ref_match.start()]
    else:
        main_body = body

    lines = main_body.split('\n')
    i = 0
    in_env = None  # equation, align, enumerate, itemize, verbatim
    env_lines = []
    pending_figure = None

    while i < len(lines):
        line = lines[i]

        # 环境追踪
        if '\\begin{equation}' in line or '\\begin{equation*}' in line:
            in_env = 'equation'
            env_lines = [line.replace('\\begin{equation}', '').replace('\\begin{equation*}', '').replace('\\end{equation}', '').replace('\\end{equation*}', '')]
            i += 1
            continue
        if '\\begin{align}' in line or '\\begin{align*}' in line:
            in_env = 'align'
            env_lines = [line.replace('\\begin{align}', '').replace('\\begin{align*}', '').replace('\\end{align}', '').replace('\\end{align*}', '')]
            i += 1
            continue
        if ('\\end{equation}' in line or '\\end{equation*}' in line or
            '\\end{align}' in line or '\\end{align*}' in line):
            if in_env in ('equation', 'align'):
                eq_text = ' '.join(env_lines).strip()
                eq_text = clean_latex_text(eq_text)
                if eq_text:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(6)
                    render_math(p, eq_text, base_size=Pt(12))
                in_env = None
                env_lines = []
            i += 1
            continue
        if in_env:
            env_lines.append(line)
            i += 1
            continue

        if '\\begin{enumerate}' in line:
            in_env = 'enumerate'
            i += 1
            continue
        if '\\end{enumerate}' in line:
            in_env = None
            i += 1
            continue

        # 跳过表格环境
        if '\\begin{table}' in line:
            while i < len(lines) and '\\end{table}' not in lines[i]:
                i += 1
            i += 1
            continue

        # 跳过算法环境
        if '\\begin{algorithm}' in line:
            while i < len(lines) and '\\end{algorithm}' not in lines[i]:
                i += 1
            i += 1
            continue

        line_stripped = line.strip()

        # 空行
        if not line_stripped:
            i += 1
            continue

        # 一级标题
        m = re.match(r'\\section\*?\{([^}]+)\}', line_stripped)
        if m:
            title = clean_latex_text(m.group(1))
            if title:
                add_heading_styled(doc, title, level=1)
            i += 1
            continue

        # 二级标题
        m = re.match(r'\\subsection\*?\{([^}]+)\}', line_stripped)
        if m:
            title = clean_latex_text(m.group(1))
            if title:
                add_heading_styled(doc, title, level=2)
            i += 1
            continue

        # 三级标题
        m = re.match(r'\\subsubsection\*?\{([^}]+)\}', line_stripped)
        if m:
            title = clean_latex_text(m.group(1))
            if title:
                add_heading_styled(doc, title, level=3)
            i += 1
            continue

        # enumerate 中的 \item
        if in_env == 'enumerate' and re.match(r'\\item', line_stripped):
            item_text = re.sub(r'\\item\s*(\[[^]]*\])?\s*', '', line_stripped)
            item_text = clean_latex_text(item_text)
            if item_text.strip():
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.74)
                p.paragraph_format.first_line_indent = Cm(-0.37)
                add_run(p, '• ' + item_text.strip(), size=Pt(12))
            i += 1
            continue

        # 图片
        if '\\includegraphics' in line_stripped:
            img_match = re.search(r'\\includegraphics(?:\[[^\]]*\])?\{([^}]+)\}', line_stripped)
            if img_match:
                img_path = img_match.group(1).strip()
                caption = ''
                for j in range(i+1, min(i+5, len(lines))):
                    cap_match = re.search(r'\\caption\{([^}]+)\}', lines[j])
                    if cap_match:
                        caption = cap_match.group(1).strip()
                        break
                img_basename = os.path.basename(img_path)
                full_img = os.path.join(FIGURES_DIR, img_basename)
                if os.path.exists(full_img):
                    add_figure(doc, img_basename, caption)
                elif os.path.exists(img_path):
                    add_figure(doc, img_path, caption)
                else:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    add_run(p, f'[图: {caption or img_basename}]', italic=True, size=Pt(10))
            i += 1
            continue

        # 跳过LaTeX格式命令行
        if re.match(r'^\\(?:centering|newpage|clearpage|pagebreak|hline|toprule|midrule|bottomrule|small|large|Large|LARGE|footnotesize|tiny|normalsize|vspace|hspace|vskip|hskip|par|noindent|indent|newline|linebreak|raggedright|setlength|renewcommand|fancyhead|fancyfoot|pagestyle|thispagestyle|numberwithin|CTEXsetup|captionsetup|hypersetup|geometry)\b', line_stripped):
            i += 1
            continue

        # 跳过纯LaTeX标签和注释
        if (re.match(r'^\\(?:label|ref|cite|bibitem|footnote|addcontentsline|renewcommand|newcommand|def|let|setcounter|stepcounter|addtocounter)\{', line_stripped) or
            line_stripped.startswith('%') or
            re.match(r'^[\{\}]', line_stripped)):
            i += 1
            continue

        # 跳过只含括号的行
        if re.match(r'^[\s\{\}]+$', line_stripped):
            i += 1
            continue

        # 正文段落（>10字符）
        if len(line_stripped) > 10:
            add_body_paragraph(doc, line_stripped)

        i += 1

    # ---- 参考文献 ----
    if ref_match:
        doc.add_page_break()
        add_heading_styled(doc, '参考文献', level=1)

        ref_section = body[ref_match.start():]
        bibitems = re.findall(r'\\bibitem\{([^}]*)\}\s*(.*?)(?=\\bibitem|\n\s*\n|\\end\{thebibliography\})', ref_section, re.DOTALL)

        if bibitems:
            for idx, (ref_key, ref_text) in enumerate(bibitems):
                text = clean_latex_text(ref_text)
                text = ' '.join(text.split())
                ref_entry = f'[{idx+1}] {text}'
                add_reference(doc, ref_entry)

    # ---- 附录 ----
    if appendix_match:
        doc.add_page_break()
        add_heading_styled(doc, '附录', level=1)

        appendix_section = body[appendix_match.start():]
        verbatim_blocks = re.findall(r'\\begin\{verbatim\}(.*?)\\end\{verbatim\}', appendix_section, re.DOTALL)
        if verbatim_blocks:
            for block in verbatim_blocks:
                add_code_block(doc, block)

    # ---- 保存 ----
    print(f'Saving: {OUTPUT_FILE}')
    doc.save(OUTPUT_FILE)
    print('Done!')
    return True


if __name__ == '__main__':
    convert_latex_to_docx()
